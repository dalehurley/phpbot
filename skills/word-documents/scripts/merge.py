#!/usr/bin/env python3
"""
Merge multiple Word documents (.docx) into a single document.

Usage:
    python3 merge.py doc1.docx doc2.docx doc3.docx -o combined.docx
    python3 merge.py doc1.docx doc2.docx -o combined.docx --page-breaks
    python3 merge.py master.docx appendix.docx -o full.docx --style-from first
"""

import sys
import os
import subprocess
import argparse


def ensure_dependencies():
    """Install python-docx and docxcompose if not available."""
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        print("Installing python-docx...", file=sys.stderr)
        result = subprocess.run(
            [sys.executable, "-m", "pip", "install", "python-docx", "-q"],
            capture_output=True, text=True,
        )
        if result.returncode != 0:
            print(f"Failed to install python-docx: {result.stderr}", file=sys.stderr)
            return False
        return True


def try_import_docxcompose():
    """Try to import docxcompose, install if needed."""
    try:
        from docxcompose.composer import Composer
        return True
    except ImportError:
        print("Installing docxcompose...", file=sys.stderr)
        result = subprocess.run(
            [sys.executable, "-m", "pip", "install", "docxcompose", "-q"],
            capture_output=True, text=True,
        )
        if result.returncode != 0:
            print(f"Note: docxcompose not available, using fallback merge.", file=sys.stderr)
            return False
        return True


def merge_with_docxcompose(file_paths, output_path, page_breaks=False):
    """Merge documents using docxcompose (preserves styles better)."""
    from docx import Document
    from docxcompose.composer import Composer

    # First document is the master
    master = Document(file_paths[0])
    composer = Composer(master)

    for path in file_paths[1:]:
        doc = Document(path)
        composer.append(doc)

    composer.save(output_path)


def merge_fallback(file_paths, output_path, page_breaks=False, style_from='first'):
    """Merge documents using python-docx directly."""
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from lxml import etree
    import copy

    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # Use first document as base
    merged = Document(file_paths[0])

    for path in file_paths[1:]:
        source = Document(path)

        if page_breaks:
            # Add a page break before appending the next document
            p_elem = etree.SubElement(merged.element.body, f'{{{W_NS}}}p')
            r_elem = etree.SubElement(p_elem, f'{{{W_NS}}}r')
            br_elem = etree.SubElement(r_elem, f'{{{W_NS}}}br')
            br_elem.set(f'{{{W_NS}}}type', 'page')

        # Copy paragraphs from source
        for para in source.paragraphs:
            # Copy the paragraph element
            new_para = copy.deepcopy(para._element)
            # Insert before the final sectPr if it exists
            sect_prs = merged.element.body.findall(f'{{{W_NS}}}sectPr')
            if sect_prs:
                sect_prs[-1].addprevious(new_para)
            else:
                merged.element.body.append(new_para)

        # Copy tables from source
        for table in source.tables:
            new_table = copy.deepcopy(table._element)
            sect_prs = merged.element.body.findall(f'{{{W_NS}}}sectPr')
            if sect_prs:
                sect_prs[-1].addprevious(new_table)
            else:
                merged.element.body.append(new_table)

        # Copy images - add relationships for any embedded images
        for rel_id, rel in source.part.rels.items():
            if "image" in rel.reltype:
                try:
                    image_part = rel.target_part
                    # Add the image part to the merged document
                    new_rel = merged.part.relate_to(
                        image_part,
                        rel.reltype,
                    )
                except Exception:
                    pass  # Skip images that can't be copied

    merged.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="Merge multiple Word documents")
    parser.add_argument("files", nargs='+',
                        help="Paths to .docx files to merge (in order)")
    parser.add_argument("-o", "--output", required=True,
                        help="Output file path")
    parser.add_argument("--page-breaks", action="store_true",
                        help="Insert page breaks between documents")
    parser.add_argument("--style-from",
                        choices=["first", "last"],
                        default="first",
                        help="Which document's styles to use (default: first)")

    args = parser.parse_args()

    # Validate input files
    file_paths = []
    for path in args.files:
        full_path = os.path.expanduser(path)
        full_path = os.path.abspath(full_path)
        if not os.path.exists(full_path):
            print(f"Error: File not found: {full_path}", file=sys.stderr)
            sys.exit(1)
        file_paths.append(full_path)

    if len(file_paths) < 2:
        print("Error: At least 2 files are required for merging", file=sys.stderr)
        sys.exit(1)

    if not ensure_dependencies():
        sys.exit(1)

    output_path = os.path.expanduser(args.output)
    output_path = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

    # Reverse file order if style-from is 'last'
    if args.style_from == 'last':
        file_paths = file_paths[::-1]

    # Try docxcompose first for better results, fallback to manual merge
    has_compose = try_import_docxcompose()

    if has_compose and not args.page_breaks:
        try:
            merge_with_docxcompose(file_paths, output_path, page_breaks=args.page_breaks)
            print(f"Merged {len(file_paths)} documents (docxcompose): {output_path}")
            file_size = os.path.getsize(output_path)
            print(f"Output size: {file_size:,} bytes")
            return
        except Exception as e:
            print(f"docxcompose failed ({e}), falling back to manual merge.", file=sys.stderr)

    try:
        merge_fallback(file_paths, output_path, page_breaks=args.page_breaks,
                       style_from=args.style_from)
        print(f"Merged {len(file_paths)} documents: {output_path}")
        file_size = os.path.getsize(output_path)
        print(f"Output size: {file_size:,} bytes")
    except Exception as e:
        print(f"Error merging documents: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
