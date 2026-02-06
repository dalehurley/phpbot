#!/usr/bin/env python3
"""
Compare two Word documents (.docx) and produce a diff report.

Outputs: Markdown diff, JSON diff, or a redline DOCX with tracked changes.

Usage:
    python3 compare.py original.docx revised.docx
    python3 compare.py original.docx revised.docx -o diff.json --format json
    python3 compare.py original.docx revised.docx -o redline.docx --format docx
    python3 compare.py original.docx revised.docx --granularity word
"""

import sys
import os
import subprocess
import argparse
import json
import difflib
from datetime import datetime, timezone


def ensure_dependencies():
    """Install python-docx if not available."""
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        print("Installing python-docx...", file=sys.stderr)
        result = subprocess.run(
            [sys.executable, "-m", "pip", "install", "python-docx", "lxml", "-q"],
            capture_output=True, text=True,
        )
        if result.returncode != 0:
            print(f"Failed to install python-docx: {result.stderr}", file=sys.stderr)
            return False
        return True


def extract_paragraphs(doc):
    """Extract paragraphs with their style information."""
    paragraphs = []
    for para in doc.paragraphs:
        paragraphs.append({
            "text": para.text,
            "style": para.style.name if para.style else "Normal",
        })
    return paragraphs


def compare_paragraphs(original_paras, revised_paras, granularity="paragraph"):
    """Compare paragraphs and return a list of differences."""
    orig_texts = [p["text"] for p in original_paras]
    rev_texts = [p["text"] for p in revised_paras]

    diffs = []
    matcher = difflib.SequenceMatcher(None, orig_texts, rev_texts)

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            for k in range(i1, i2):
                diffs.append({
                    "type": "equal",
                    "text": orig_texts[k],
                    "style": original_paras[k]["style"],
                })
        elif tag == 'replace':
            # For word-level granularity, diff the individual changed paragraphs
            if granularity == "word":
                for oi in range(i1, i2):
                    for ri in range(j1, j2):
                        word_diffs = compare_words(orig_texts[oi], rev_texts[ri])
                        diffs.append({
                            "type": "modified",
                            "original_text": orig_texts[oi],
                            "revised_text": rev_texts[ri],
                            "style": original_paras[oi]["style"],
                            "word_diffs": word_diffs,
                        })
            else:
                for k in range(i1, i2):
                    diffs.append({
                        "type": "deleted",
                        "text": orig_texts[k],
                        "style": original_paras[k]["style"],
                    })
                for k in range(j1, j2):
                    diffs.append({
                        "type": "inserted",
                        "text": rev_texts[k],
                        "style": revised_paras[k]["style"],
                    })
        elif tag == 'delete':
            for k in range(i1, i2):
                diffs.append({
                    "type": "deleted",
                    "text": orig_texts[k],
                    "style": original_paras[k]["style"],
                })
        elif tag == 'insert':
            for k in range(j1, j2):
                diffs.append({
                    "type": "inserted",
                    "text": rev_texts[k],
                    "style": revised_paras[k]["style"],
                })

    return diffs


def compare_words(original, revised):
    """Perform word-level comparison between two strings."""
    orig_words = original.split()
    rev_words = revised.split()

    word_diffs = []
    matcher = difflib.SequenceMatcher(None, orig_words, rev_words)

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            word_diffs.append({
                "type": "equal",
                "text": " ".join(orig_words[i1:i2]),
            })
        elif tag == 'replace':
            word_diffs.append({
                "type": "deleted",
                "text": " ".join(orig_words[i1:i2]),
            })
            word_diffs.append({
                "type": "inserted",
                "text": " ".join(rev_words[j1:j2]),
            })
        elif tag == 'delete':
            word_diffs.append({
                "type": "deleted",
                "text": " ".join(orig_words[i1:i2]),
            })
        elif tag == 'insert':
            word_diffs.append({
                "type": "inserted",
                "text": " ".join(rev_words[j1:j2]),
            })

    return word_diffs


def format_markdown(diffs, original_path, revised_path):
    """Format diff results as Markdown."""
    lines = []
    lines.append(f"# Document Comparison")
    lines.append(f"")
    lines.append(f"**Original:** {os.path.basename(original_path)}")
    lines.append(f"**Revised:** {os.path.basename(revised_path)}")
    lines.append(f"**Date:** {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    lines.append(f"")

    # Statistics
    insertions = sum(1 for d in diffs if d['type'] == 'inserted')
    deletions = sum(1 for d in diffs if d['type'] == 'deleted')
    modifications = sum(1 for d in diffs if d['type'] == 'modified')
    unchanged = sum(1 for d in diffs if d['type'] == 'equal')

    lines.append(f"## Summary")
    lines.append(f"")
    lines.append(f"- Unchanged paragraphs: {unchanged}")
    lines.append(f"- Inserted paragraphs: {insertions}")
    lines.append(f"- Deleted paragraphs: {deletions}")
    if modifications:
        lines.append(f"- Modified paragraphs: {modifications}")
    lines.append(f"")

    lines.append(f"## Changes")
    lines.append(f"")

    for diff in diffs:
        if diff['type'] == 'equal':
            text = diff['text']
            if len(text) > 80:
                text = text[:77] + "..."
            lines.append(f"  {text}")
        elif diff['type'] == 'inserted':
            lines.append(f"+ {diff['text']}")
        elif diff['type'] == 'deleted':
            lines.append(f"- {diff['text']}")
        elif diff['type'] == 'modified':
            lines.append(f"~ Original: {diff['original_text']}")
            lines.append(f"~ Revised:  {diff['revised_text']}")
            if 'word_diffs' in diff:
                parts = []
                for wd in diff['word_diffs']:
                    if wd['type'] == 'equal':
                        parts.append(wd['text'])
                    elif wd['type'] == 'deleted':
                        parts.append(f"[-{wd['text']}-]")
                    elif wd['type'] == 'inserted':
                        parts.append(f"[+{wd['text']}+]")
                lines.append(f"  Detail: {' '.join(parts)}")

    return "\n".join(lines)


def format_json(diffs, original_path, revised_path):
    """Format diff results as JSON."""
    insertions = sum(1 for d in diffs if d['type'] == 'inserted')
    deletions = sum(1 for d in diffs if d['type'] == 'deleted')
    modifications = sum(1 for d in diffs if d['type'] == 'modified')
    unchanged = sum(1 for d in diffs if d['type'] == 'equal')

    result = {
        "original": os.path.basename(original_path),
        "revised": os.path.basename(revised_path),
        "date": datetime.now().isoformat(),
        "summary": {
            "unchanged": unchanged,
            "inserted": insertions,
            "deleted": deletions,
            "modified": modifications,
        },
        "diffs": diffs,
    }
    return json.dumps(result, indent=2, ensure_ascii=False)


def create_redline_docx(original_doc, diffs, output_path):
    """Create a DOCX with tracked changes showing the differences."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from lxml import etree

    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    change_id = 0

    # Start from a copy of the original
    redline = Document(original_doc)

    # Clear existing content
    body = redline.element.body
    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'sectPr':
            continue  # Keep section properties
        body.remove(child)

    for diff in diffs:
        if diff['type'] == 'equal':
            # Add as normal paragraph
            p = etree.SubElement(body, f'{{{W_NS}}}p')
            if diff.get('style') and diff['style'] != 'Normal':
                ppr = etree.SubElement(p, f'{{{W_NS}}}pPr')
                pstyle = etree.SubElement(ppr, f'{{{W_NS}}}pStyle')
                pstyle.set(f'{{{W_NS}}}val', diff['style'].replace(' ', ''))
            r = etree.SubElement(p, f'{{{W_NS}}}r')
            t = etree.SubElement(r, f'{{{W_NS}}}t')
            t.text = diff['text']
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

        elif diff['type'] == 'inserted':
            # Add as insertion
            p = etree.SubElement(body, f'{{{W_NS}}}p')
            if diff.get('style') and diff['style'] != 'Normal':
                ppr = etree.SubElement(p, f'{{{W_NS}}}pPr')
                pstyle = etree.SubElement(ppr, f'{{{W_NS}}}pStyle')
                pstyle.set(f'{{{W_NS}}}val', diff['style'].replace(' ', ''))
            ins = etree.SubElement(p, f'{{{W_NS}}}ins')
            ins.set(f'{{{W_NS}}}id', str(change_id))
            ins.set(f'{{{W_NS}}}author', 'Document Compare')
            ins.set(f'{{{W_NS}}}date', now)
            change_id += 1
            r = etree.SubElement(ins, f'{{{W_NS}}}r')
            t = etree.SubElement(r, f'{{{W_NS}}}t')
            t.text = diff['text']
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

        elif diff['type'] == 'deleted':
            # Add as deletion
            p = etree.SubElement(body, f'{{{W_NS}}}p')
            if diff.get('style') and diff['style'] != 'Normal':
                ppr = etree.SubElement(p, f'{{{W_NS}}}pPr')
                pstyle = etree.SubElement(ppr, f'{{{W_NS}}}pStyle')
                pstyle.set(f'{{{W_NS}}}val', diff['style'].replace(' ', ''))
            dele = etree.SubElement(p, f'{{{W_NS}}}del')
            dele.set(f'{{{W_NS}}}id', str(change_id))
            dele.set(f'{{{W_NS}}}author', 'Document Compare')
            dele.set(f'{{{W_NS}}}date', now)
            change_id += 1
            r = etree.SubElement(dele, f'{{{W_NS}}}r')
            dt = etree.SubElement(r, f'{{{W_NS}}}delText')
            dt.text = diff['text']
            dt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

        elif diff['type'] == 'modified':
            # Add as deletion + insertion
            p = etree.SubElement(body, f'{{{W_NS}}}p')
            if diff.get('style') and diff['style'] != 'Normal':
                ppr = etree.SubElement(p, f'{{{W_NS}}}pPr')
                pstyle = etree.SubElement(ppr, f'{{{W_NS}}}pStyle')
                pstyle.set(f'{{{W_NS}}}val', diff['style'].replace(' ', ''))

            dele = etree.SubElement(p, f'{{{W_NS}}}del')
            dele.set(f'{{{W_NS}}}id', str(change_id))
            dele.set(f'{{{W_NS}}}author', 'Document Compare')
            dele.set(f'{{{W_NS}}}date', now)
            change_id += 1
            r = etree.SubElement(dele, f'{{{W_NS}}}r')
            dt = etree.SubElement(r, f'{{{W_NS}}}delText')
            dt.text = diff['original_text']
            dt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

            ins = etree.SubElement(p, f'{{{W_NS}}}ins')
            ins.set(f'{{{W_NS}}}id', str(change_id))
            ins.set(f'{{{W_NS}}}author', 'Document Compare')
            ins.set(f'{{{W_NS}}}date', now)
            change_id += 1
            r = etree.SubElement(ins, f'{{{W_NS}}}r')
            t = etree.SubElement(r, f'{{{W_NS}}}t')
            t.text = diff['revised_text']
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    redline.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="Compare two Word documents")
    parser.add_argument("original", help="Path to original .docx file")
    parser.add_argument("revised", help="Path to revised .docx file")
    parser.add_argument("-o", "--output", default=None,
                        help="Output file path (default: stdout)")
    parser.add_argument("--format", choices=["markdown", "json", "docx"], default="markdown",
                        help="Output format (default: markdown)")
    parser.add_argument("--granularity", choices=["paragraph", "word"], default="paragraph",
                        help="Comparison granularity (default: paragraph)")

    args = parser.parse_args()

    original_path = os.path.expanduser(args.original)
    original_path = os.path.abspath(original_path)
    revised_path = os.path.expanduser(args.revised)
    revised_path = os.path.abspath(revised_path)

    for path in [original_path, revised_path]:
        if not os.path.exists(path):
            print(f"Error: File not found: {path}", file=sys.stderr)
            sys.exit(1)

    if not ensure_dependencies():
        sys.exit(1)

    from docx import Document

    try:
        original_doc = Document(original_path)
        revised_doc = Document(revised_path)
    except Exception as e:
        print(f"Error: Could not open document: {e}", file=sys.stderr)
        sys.exit(1)

    original_paras = extract_paragraphs(original_doc)
    revised_paras = extract_paragraphs(revised_doc)

    diffs = compare_paragraphs(original_paras, revised_paras, granularity=args.granularity)

    if args.format == "docx":
        if not args.output:
            print("Error: --output is required for DOCX format", file=sys.stderr)
            sys.exit(1)
        output_path = os.path.expanduser(args.output)
        output_path = os.path.abspath(output_path)
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        create_redline_docx(original_path, diffs, output_path)
        print(f"Redline document saved to: {output_path}", file=sys.stderr)
    elif args.format == "json":
        result = format_json(diffs, original_path, revised_path)
        if args.output:
            output_path = os.path.expanduser(args.output)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(result)
            print(f"JSON diff saved to: {output_path}", file=sys.stderr)
        else:
            print(result)
    else:
        result = format_markdown(diffs, original_path, revised_path)
        if args.output:
            output_path = os.path.expanduser(args.output)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(result)
            print(f"Markdown diff saved to: {output_path}", file=sys.stderr)
        else:
            print(result)


if __name__ == "__main__":
    main()
