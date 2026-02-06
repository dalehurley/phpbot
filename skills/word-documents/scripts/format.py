#!/usr/bin/env python3
"""
Apply formatting standards to a Word document (.docx).

Normalizes styles, fonts, spacing, and heading hierarchy to named standards
(legal, academic, corporate, regulatory) or a custom JSON standard file.

Usage:
    python3 format.py document.docx -s legal -o formatted.docx
    python3 format.py document.docx -s academic -o formatted.docx
    python3 format.py document.docx --custom-standard my_standard.json -o formatted.docx
    python3 format.py document.docx --fix-headings -o formatted.docx
    python3 format.py document.docx --clean-direct-formatting -o formatted.docx
"""

import sys
import os
import subprocess
import argparse
import json


def ensure_dependencies():
    """Install python-docx if not available."""
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        print("Installing python-docx...", file=sys.stderr)
        result = subprocess.run(
            [sys.executable, "-m", "pip", "install", "python-docx", "-q"],
            capture_output=True, text=True,
        )
        if result.returncode != 0:
            print(f"Failed to install python-docx: {result.stderr}", file=sys.stderr)
            return False
        return True


# Built-in formatting standards
STANDARDS = {
    "legal": {
        "name": "Legal / Court Filing",
        "body": {
            "font_name": "Times New Roman",
            "font_size_pt": 12,
            "line_spacing": 2.0,
            "space_before_pt": 0,
            "space_after_pt": 0,
            "alignment": "justify",
            "first_line_indent_inches": 0.5,
        },
        "heading1": {
            "font_name": "Times New Roman",
            "font_size_pt": 14,
            "bold": True,
            "alignment": "center",
            "space_before_pt": 24,
            "space_after_pt": 12,
            "all_caps": True,
        },
        "heading2": {
            "font_name": "Times New Roman",
            "font_size_pt": 13,
            "bold": True,
            "alignment": "left",
            "space_before_pt": 18,
            "space_after_pt": 6,
        },
        "heading3": {
            "font_name": "Times New Roman",
            "font_size_pt": 12,
            "bold": True,
            "italic": True,
            "alignment": "left",
            "space_before_pt": 12,
            "space_after_pt": 6,
        },
        "page": {
            "margin_top_inches": 1.0,
            "margin_bottom_inches": 1.0,
            "margin_left_inches": 1.0,
            "margin_right_inches": 1.0,
        },
    },
    "academic": {
        "name": "Academic (APA-style)",
        "body": {
            "font_name": "Times New Roman",
            "font_size_pt": 12,
            "line_spacing": 2.0,
            "space_before_pt": 0,
            "space_after_pt": 0,
            "alignment": "left",
            "first_line_indent_inches": 0.5,
        },
        "heading1": {
            "font_name": "Times New Roman",
            "font_size_pt": 12,
            "bold": True,
            "alignment": "center",
            "space_before_pt": 24,
            "space_after_pt": 12,
        },
        "heading2": {
            "font_name": "Times New Roman",
            "font_size_pt": 12,
            "bold": True,
            "alignment": "left",
            "space_before_pt": 18,
            "space_after_pt": 6,
        },
        "heading3": {
            "font_name": "Times New Roman",
            "font_size_pt": 12,
            "bold": True,
            "italic": True,
            "alignment": "left",
            "space_before_pt": 12,
            "space_after_pt": 6,
        },
        "page": {
            "margin_top_inches": 1.0,
            "margin_bottom_inches": 1.0,
            "margin_left_inches": 1.0,
            "margin_right_inches": 1.0,
        },
    },
    "corporate": {
        "name": "Corporate / Business",
        "body": {
            "font_name": "Calibri",
            "font_size_pt": 11,
            "line_spacing": 1.15,
            "space_before_pt": 0,
            "space_after_pt": 8,
            "alignment": "left",
        },
        "heading1": {
            "font_name": "Calibri",
            "font_size_pt": 16,
            "bold": True,
            "color": "2E74B5",
            "alignment": "left",
            "space_before_pt": 24,
            "space_after_pt": 6,
        },
        "heading2": {
            "font_name": "Calibri",
            "font_size_pt": 13,
            "bold": True,
            "color": "2E74B5",
            "alignment": "left",
            "space_before_pt": 18,
            "space_after_pt": 4,
        },
        "heading3": {
            "font_name": "Calibri",
            "font_size_pt": 12,
            "bold": True,
            "color": "1F4D78",
            "alignment": "left",
            "space_before_pt": 12,
            "space_after_pt": 4,
        },
        "page": {
            "margin_top_inches": 1.0,
            "margin_bottom_inches": 1.0,
            "margin_left_inches": 1.0,
            "margin_right_inches": 1.0,
        },
    },
    "regulatory": {
        "name": "Regulatory / Compliance",
        "body": {
            "font_name": "Arial",
            "font_size_pt": 11,
            "line_spacing": 1.5,
            "space_before_pt": 0,
            "space_after_pt": 6,
            "alignment": "justify",
        },
        "heading1": {
            "font_name": "Arial",
            "font_size_pt": 14,
            "bold": True,
            "alignment": "left",
            "space_before_pt": 24,
            "space_after_pt": 12,
            "all_caps": True,
        },
        "heading2": {
            "font_name": "Arial",
            "font_size_pt": 12,
            "bold": True,
            "alignment": "left",
            "space_before_pt": 18,
            "space_after_pt": 6,
        },
        "heading3": {
            "font_name": "Arial",
            "font_size_pt": 11,
            "bold": True,
            "alignment": "left",
            "space_before_pt": 12,
            "space_after_pt": 6,
        },
        "page": {
            "margin_top_inches": 1.0,
            "margin_bottom_inches": 1.0,
            "margin_left_inches": 1.25,
            "margin_right_inches": 1.25,
        },
    },
    "accessibility": {
        "name": "Accessible Document",
        "body": {
            "font_name": "Arial",
            "font_size_pt": 12,
            "line_spacing": 1.5,
            "space_before_pt": 0,
            "space_after_pt": 8,
            "alignment": "left",
        },
        "heading1": {
            "font_name": "Arial",
            "font_size_pt": 18,
            "bold": True,
            "alignment": "left",
            "space_before_pt": 36,
            "space_after_pt": 12,
        },
        "heading2": {
            "font_name": "Arial",
            "font_size_pt": 16,
            "bold": True,
            "alignment": "left",
            "space_before_pt": 24,
            "space_after_pt": 8,
        },
        "heading3": {
            "font_name": "Arial",
            "font_size_pt": 14,
            "bold": True,
            "alignment": "left",
            "space_before_pt": 18,
            "space_after_pt": 6,
        },
        "page": {
            "margin_top_inches": 1.0,
            "margin_bottom_inches": 1.0,
            "margin_left_inches": 1.0,
            "margin_right_inches": 1.0,
        },
    },
}


def inches_to_emu(inches):
    """Convert inches to EMU (English Metric Units)."""
    return int(inches * 914400)


def inches_to_twips(inches):
    """Convert inches to twips (1/1440 inch)."""
    return int(inches * 1440)


def pt_to_half_pt(pt):
    """Convert points to half-points (python-docx uses Pt objects)."""
    return int(pt * 2)


def apply_run_format(run, fmt):
    """Apply formatting to a run."""
    from docx.shared import Pt, RGBColor

    if 'font_name' in fmt:
        run.font.name = fmt['font_name']
    if 'font_size_pt' in fmt:
        run.font.size = Pt(fmt['font_size_pt'])
    if 'bold' in fmt:
        run.font.bold = fmt['bold']
    if 'italic' in fmt:
        run.font.italic = fmt['italic']
    if 'color' in fmt:
        run.font.color.rgb = RGBColor.from_string(fmt['color'])
    if 'all_caps' in fmt:
        run.font.all_caps = fmt['all_caps']


def apply_paragraph_format(para, fmt):
    """Apply formatting to a paragraph."""
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    pf = para.paragraph_format

    alignment_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    if 'alignment' in fmt:
        pf.alignment = alignment_map.get(fmt['alignment'], WD_ALIGN_PARAGRAPH.LEFT)
    if 'space_before_pt' in fmt:
        pf.space_before = Pt(fmt['space_before_pt'])
    if 'space_after_pt' in fmt:
        pf.space_after = Pt(fmt['space_after_pt'])
    if 'line_spacing' in fmt:
        pf.line_spacing = fmt['line_spacing']
    if 'first_line_indent_inches' in fmt:
        pf.first_line_indent = Inches(fmt['first_line_indent_inches'])

    # Apply run formatting to all runs in the paragraph
    for run in para.runs:
        apply_run_format(run, fmt)


def apply_page_format(doc, page_fmt):
    """Apply page layout formatting to all sections."""
    from docx.shared import Inches

    for section in doc.sections:
        if 'margin_top_inches' in page_fmt:
            section.top_margin = Inches(page_fmt['margin_top_inches'])
        if 'margin_bottom_inches' in page_fmt:
            section.bottom_margin = Inches(page_fmt['margin_bottom_inches'])
        if 'margin_left_inches' in page_fmt:
            section.left_margin = Inches(page_fmt['margin_left_inches'])
        if 'margin_right_inches' in page_fmt:
            section.right_margin = Inches(page_fmt['margin_right_inches'])


def apply_standard(doc, standard):
    """Apply a complete formatting standard to the document."""
    applied = {"body": 0, "heading1": 0, "heading2": 0, "heading3": 0}

    # Apply page formatting
    if 'page' in standard:
        apply_page_format(doc, standard['page'])

    heading_map = {
        'Heading 1': 'heading1',
        'Heading1': 'heading1',
        'Heading 2': 'heading2',
        'Heading2': 'heading2',
        'Heading 3': 'heading3',
        'Heading3': 'heading3',
    }

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else 'Normal'

        if style_name in heading_map:
            key = heading_map[style_name]
            if key in standard:
                apply_paragraph_format(para, standard[key])
                applied[key] = applied.get(key, 0) + 1
        elif style_name in ('Normal', 'Body Text', 'Body', 'Default'):
            if 'body' in standard:
                apply_paragraph_format(para, standard['body'])
                applied['body'] += 1
        elif style_name.startswith('List') or style_name.startswith('Bullet'):
            # Apply body font to list items but skip indent changes
            if 'body' in standard:
                body_fmt = {k: v for k, v in standard['body'].items()
                            if k not in ('first_line_indent_inches', 'alignment')}
                apply_paragraph_format(para, body_fmt)
                applied['body'] += 1

    return applied


def fix_heading_hierarchy(doc):
    """Fix heading hierarchy to ensure proper nesting (H1 > H2 > H3)."""
    from docx.enum.style import WD_STYLE_TYPE

    heading_paras = []
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''
        if style_name.startswith('Heading'):
            try:
                level = int(style_name.replace('Heading ', '').replace('Heading', '').strip())
                heading_paras.append((para, level))
            except ValueError:
                pass

    if not heading_paras:
        print("No headings found in the document.")
        return 0

    # Check and fix hierarchy
    fixes = 0
    prev_level = 0

    for para, level in heading_paras:
        if level > prev_level + 1:
            # Gap in hierarchy (e.g., H1 followed by H3)
            correct_level = prev_level + 1
            new_style_name = f'Heading {correct_level}'
            try:
                para.style = doc.styles[new_style_name]
                fixes += 1
                print(f"  Fixed: \"{para.text[:50]}\" from Heading {level} -> Heading {correct_level}")
            except KeyError:
                print(f"  Warning: Style '{new_style_name}' not found", file=sys.stderr)
        prev_level = level if level <= prev_level + 1 else prev_level + 1

    return fixes


def clean_direct_formatting(doc):
    """Remove direct formatting from runs, relying on styles instead."""
    cleaned = 0

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else 'Normal'

        for run in para.runs:
            # Only clean runs that have direct formatting overriding the style
            # Keep bold/italic if the style is a heading (they're intentional)
            if not style_name.startswith('Heading'):
                had_changes = False
                if run.font.name is not None:
                    run.font.name = None
                    had_changes = True
                if run.font.size is not None:
                    run.font.size = None
                    had_changes = True
                if had_changes:
                    cleaned += 1

    return cleaned


def load_custom_standard(path):
    """Load a custom standard from a JSON file."""
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


def main():
    parser = argparse.ArgumentParser(description="Apply formatting standards to a Word document")
    parser.add_argument("file_path", help="Path to .docx file")
    parser.add_argument("-s", "--standard",
                        choices=list(STANDARDS.keys()),
                        help="Named formatting standard to apply")
    parser.add_argument("--custom-standard",
                        help="Path to custom standard JSON file")
    parser.add_argument("--fix-headings", action="store_true",
                        help="Fix heading hierarchy only")
    parser.add_argument("--clean-direct-formatting", action="store_true",
                        help="Clean up direct formatting only")
    parser.add_argument("-o", "--output", required=True,
                        help="Output file path")

    args = parser.parse_args()

    if not args.standard and not args.custom_standard and not args.fix_headings and not args.clean_direct_formatting:
        print("Error: Specify --standard, --custom-standard, --fix-headings, or --clean-direct-formatting",
              file=sys.stderr)
        sys.exit(1)

    file_path = os.path.expanduser(args.file_path)
    file_path = os.path.abspath(file_path)

    if not os.path.exists(file_path):
        print(f"Error: File not found: {file_path}", file=sys.stderr)
        sys.exit(1)

    if not ensure_dependencies():
        sys.exit(1)

    from docx import Document

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"Error: Could not open document: {e}", file=sys.stderr)
        sys.exit(1)

    output_path = os.path.expanduser(args.output)
    output_path = os.path.abspath(output_path)

    if args.fix_headings:
        fixes = fix_heading_hierarchy(doc)
        print(f"Fixed {fixes} heading(s)")

    if args.clean_direct_formatting:
        cleaned = clean_direct_formatting(doc)
        print(f"Cleaned direct formatting from {cleaned} run(s)")

    if args.standard:
        standard = STANDARDS[args.standard]
        print(f"Applying standard: {standard['name']}")
        applied = apply_standard(doc, standard)
        for key, count in applied.items():
            if count > 0:
                print(f"  {key}: {count} paragraph(s)")

    if args.custom_standard:
        custom_path = os.path.expanduser(args.custom_standard)
        if not os.path.exists(custom_path):
            print(f"Error: Custom standard file not found: {custom_path}", file=sys.stderr)
            sys.exit(1)
        standard = load_custom_standard(custom_path)
        print(f"Applying custom standard: {standard.get('name', 'Custom')}")
        applied = apply_standard(doc, standard)
        for key, count in applied.items():
            if count > 0:
                print(f"  {key}: {count} paragraph(s)")

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    print(f"\nSaved to: {output_path}")


if __name__ == "__main__":
    main()
