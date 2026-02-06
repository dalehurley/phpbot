#!/usr/bin/env python3
"""
Analyze a Word document (.docx) and extract structured information.

Extracts: text, metadata, styles, comments, tracked changes, structure (headings),
word/paragraph/table/image counts.

Usage:
    python3 analyze.py document.docx [--mode full|text|metadata|comments|changes|structure|styles] [-o output.json]
"""

import sys
import os
import subprocess
import argparse
import json
from datetime import datetime


def ensure_dependencies():
    """Install python-docx and lxml if not available."""
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        print("Installing python-docx...", file=sys.stderr)
        result = subprocess.run(
            [sys.executable, "-m", "pip", "install", "python-docx", "lxml", "-q"],
            capture_output=True, text=True,
        )
        if result.returncode != 0:
            print(f"Failed to install python-docx: {result.stderr}", file=sys.stderr)
            return False
        return True


def extract_metadata(doc):
    """Extract document metadata."""
    props = doc.core_properties
    meta = {}
    for attr in ['author', 'title', 'subject', 'keywords', 'category',
                 'comments', 'last_modified_by', 'revision', 'version',
                 'content_status', 'identifier', 'language']:
        val = getattr(props, attr, None)
        if val is not None:
            meta[attr] = str(val) if val else None
    for attr in ['created', 'modified', 'last_printed']:
        val = getattr(props, attr, None)
        if val is not None:
            meta[attr] = val.isoformat() if isinstance(val, datetime) else str(val)
    return meta


def extract_text(doc):
    """Extract full text from document."""
    paragraphs = []
    for para in doc.paragraphs:
        paragraphs.append(para.text)
    return "\n".join(paragraphs)


def extract_structure(doc):
    """Extract heading hierarchy."""
    headings = []
    for para in doc.paragraphs:
        if para.style and para.style.name and para.style.name.startswith('Heading'):
            try:
                level = int(para.style.name.replace('Heading ', '').replace('Heading', '').strip())
            except ValueError:
                level = 0
            headings.append({
                "level": level,
                "text": para.text.strip(),
            })
    return headings


def extract_styles(doc):
    """Extract all styles in use."""
    paragraph_styles = set()
    character_styles = set()
    for para in doc.paragraphs:
        if para.style:
            paragraph_styles.add(para.style.name)
        for run in para.runs:
            if run.style and run.style.name != 'Default Paragraph Font':
                character_styles.add(run.style.name)

    all_styles = []
    for style in doc.styles:
        all_styles.append({
            "name": style.name,
            "type": str(style.type),
            "builtin": style.builtin,
            "in_use": style.name in paragraph_styles or style.name in character_styles,
        })

    return {
        "paragraph_styles_in_use": sorted(paragraph_styles),
        "character_styles_in_use": sorted(character_styles),
        "all_styles": all_styles,
    }


def extract_comments(doc):
    """Extract comments using lxml to parse the underlying XML."""
    from lxml import etree

    comments = []
    nsmap = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'w15': 'http://schemas.microsoft.com/office/word/2012/wordml',
        'w16cid': 'http://schemas.microsoft.com/office/word/2016/wordml/cid',
    }

    # Access the comments part from the package
    try:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        comments_part = None
        for rel in doc.part.rels.values():
            if 'comments' in rel.reltype and 'Extended' not in rel.reltype:
                comments_part = rel.target_part
                break

        if comments_part is None:
            return comments

        root = etree.fromstring(comments_part.blob)
        for comment_elem in root.findall('.//w:comment', nsmap):
            comment_id = comment_elem.get(f'{{{nsmap["w"]}}}id')
            author = comment_elem.get(f'{{{nsmap["w"]}}}author', 'Unknown')
            date = comment_elem.get(f'{{{nsmap["w"]}}}date', '')
            initials = comment_elem.get(f'{{{nsmap["w"]}}}initials', '')

            # Extract comment text from paragraphs
            text_parts = []
            for p in comment_elem.findall('.//w:p', nsmap):
                for r in p.findall('.//w:r', nsmap):
                    for t in r.findall('.//w:t', nsmap):
                        if t.text:
                            text_parts.append(t.text)

            comments.append({
                "id": comment_id,
                "author": author,
                "date": date,
                "initials": initials,
                "text": " ".join(text_parts),
            })
    except Exception as e:
        comments.append({"error": f"Could not extract comments: {str(e)}"})

    return comments


def extract_tracked_changes(doc):
    """Extract tracked changes using lxml to parse the underlying XML."""
    from lxml import etree

    changes = []
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    body = doc.element.body

    # Find insertions
    for ins in body.iter(f'{{{nsmap["w"]}}}ins'):
        change_id = ins.get(f'{{{nsmap["w"]}}}id', '')
        author = ins.get(f'{{{nsmap["w"]}}}author', 'Unknown')
        date = ins.get(f'{{{nsmap["w"]}}}date', '')

        text_parts = []
        for t in ins.iter(f'{{{nsmap["w"]}}}t'):
            if t.text:
                text_parts.append(t.text)

        if text_parts:
            changes.append({
                "type": "insertion",
                "id": change_id,
                "author": author,
                "date": date,
                "text": "".join(text_parts),
            })

    # Find deletions
    for dele in body.iter(f'{{{nsmap["w"]}}}del'):
        change_id = dele.get(f'{{{nsmap["w"]}}}id', '')
        author = dele.get(f'{{{nsmap["w"]}}}author', 'Unknown')
        date = dele.get(f'{{{nsmap["w"]}}}date', '')

        text_parts = []
        for dt in dele.iter(f'{{{nsmap["w"]}}}delText'):
            if dt.text:
                text_parts.append(dt.text)

        if text_parts:
            changes.append({
                "type": "deletion",
                "id": change_id,
                "author": author,
                "date": date,
                "text": "".join(text_parts),
            })

    # Find format changes
    for rpc in body.iter(f'{{{nsmap["w"]}}}rPrChange'):
        change_id = rpc.get(f'{{{nsmap["w"]}}}id', '')
        author = rpc.get(f'{{{nsmap["w"]}}}author', 'Unknown')
        date = rpc.get(f'{{{nsmap["w"]}}}date', '')

        # Get the parent run's text
        parent_run = rpc.getparent()
        if parent_run is not None:
            parent_r = parent_run.getparent()
            if parent_r is not None:
                text_parts = []
                for t in parent_r.iter(f'{{{nsmap["w"]}}}t'):
                    if t.text:
                        text_parts.append(t.text)

                if text_parts:
                    changes.append({
                        "type": "format_change",
                        "id": change_id,
                        "author": author,
                        "date": date,
                        "text": "".join(text_parts),
                    })

    return changes


def count_statistics(doc):
    """Count paragraphs, words, tables, images."""
    total_text = extract_text(doc)
    words = len(total_text.split())
    characters = len(total_text)

    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1

    return {
        "paragraphs": len(doc.paragraphs),
        "words": words,
        "characters": characters,
        "tables": len(doc.tables),
        "sections": len(doc.sections),
        "images": image_count,
    }


def full_analysis(doc):
    """Run all analysis modes and return combined result."""
    return {
        "metadata": extract_metadata(doc),
        "statistics": count_statistics(doc),
        "structure": extract_structure(doc),
        "styles": extract_styles(doc),
        "comments": extract_comments(doc),
        "tracked_changes": extract_tracked_changes(doc),
    }


def main():
    parser = argparse.ArgumentParser(
        description="Analyze a Word document (.docx)"
    )
    parser.add_argument("file_path", help="Path to the .docx file")
    parser.add_argument(
        "--mode", "-m",
        choices=["full", "text", "metadata", "comments", "changes", "structure", "styles"],
        default="full",
        help="Analysis mode (default: full)"
    )
    parser.add_argument(
        "-o", "--output",
        help="Output file path for JSON (default: stdout)",
        default=None
    )

    args = parser.parse_args()

    file_path = os.path.expanduser(args.file_path)
    file_path = os.path.abspath(file_path)

    if not os.path.exists(file_path):
        print(f"Error: File not found: {file_path}", file=sys.stderr)
        sys.exit(1)

    if not ensure_dependencies():
        sys.exit(1)

    from docx import Document

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"Error: Could not open document: {e}", file=sys.stderr)
        sys.exit(1)

    mode_map = {
        "full": lambda: full_analysis(doc),
        "text": lambda: {"text": extract_text(doc)},
        "metadata": lambda: {"metadata": extract_metadata(doc)},
        "comments": lambda: {"comments": extract_comments(doc)},
        "changes": lambda: {"tracked_changes": extract_tracked_changes(doc)},
        "structure": lambda: {"structure": extract_structure(doc)},
        "styles": lambda: {"styles": extract_styles(doc)},
    }

    result = mode_map[args.mode]()

    output_json = json.dumps(result, indent=2, default=str, ensure_ascii=False)

    if args.output:
        output_path = os.path.expanduser(args.output)
        output_path = os.path.abspath(output_path)
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(output_json)
        print(f"Analysis written to: {output_path}", file=sys.stderr)
    else:
        print(output_json)


if __name__ == "__main__":
    main()
