#!/usr/bin/env python3
"""
Create a new Word document (.docx) from text, Markdown, or a JSON specification.

Usage:
    python3 create.py -i content.md -o report.docx
    python3 create.py -i notes.txt -o document.docx
    python3 create.py -i spec.json -o report.docx
    python3 create.py -i content.md -o report.docx --template template.docx
    python3 create.py -i content.md -o report.docx --header "Confidential" --footer "Page {page}"
    python3 create.py --text "Hello World" -o hello.docx
"""

import sys
import os
import subprocess
import argparse
import json
import re


def ensure_dependencies():
    """Install python-docx if not available."""
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        print("Installing python-docx...", file=sys.stderr)
        result = subprocess.run(
            [sys.executable, "-m", "pip", "install", "python-docx", "-q"],
            capture_output=True, text=True,
        )
        if result.returncode != 0:
            print(f"Failed to install python-docx: {result.stderr}", file=sys.stderr)
            return False
        return True


def detect_input_type(file_path):
    """Detect input type from file extension."""
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    if ext == '.json':
        return 'json'
    elif ext in ('.md', '.markdown'):
        return 'markdown'
    else:
        return 'text'


def create_from_text(text, doc, title=None):
    """Create document content from plain text."""
    from docx.shared import Pt

    if title:
        doc.add_heading(title, level=0)

    for line in text.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph('')


def create_from_markdown(text, doc):
    """Create document content from Markdown text."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lines = text.split('\n')
    i = 0
    in_code_block = False
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                i += 1
                continue

        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            i += 1
            continue

        # Table detection
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            # Parse table row
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            # Skip separator rows (---|---|---)
            if not all(c.replace('-', '').replace(':', '').strip() == '' for c in cells):
                table_rows.append(cells)
            i += 1
            # Check if next line continues the table
            if i >= len(lines) or '|' not in lines[i]:
                # End of table, create it
                if table_rows:
                    _create_table(doc, table_rows)
                in_table = False
                table_rows = []
            continue

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text_content = heading_match.group(2).strip()
            doc.add_heading(text_content, level=min(level, 9))
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^[-*_]{3,}\s*$', line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            # Add a thin line by using a bottom border (simplified)
            run = p.add_run('─' * 50)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            i += 1
            continue

        # Bullet lists
        bullet_match = re.match(r'^(\s*)[-*+]\s+(.+)$', line)
        if bullet_match:
            indent_level = len(bullet_match.group(1)) // 2
            text_content = bullet_match.group(2)
            p = doc.add_paragraph(style='List Bullet')
            _add_inline_formatting(p, text_content)
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r'^(\s*)\d+[.)]\s+(.+)$', line)
        if num_match:
            text_content = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            _add_inline_formatting(p, text_content)
            i += 1
            continue

        # Blockquote
        if line.startswith('>'):
            text_content = line.lstrip('>').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(36)
            run = p.add_run(text_content)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_inline_formatting(p, line)
        i += 1


def _add_inline_formatting(paragraph, text):
    """Parse inline Markdown formatting and add runs to paragraph."""
    from docx.shared import Pt, RGBColor

    # Simple inline formatting parser
    # Handles: **bold**, *italic*, `code`, ***bold-italic***
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)'   # bold-italic
        r'|(\*\*(.+?)\*\*)'       # bold
        r'|(\*(.+?)\*)'           # italic
        r'|(`(.+?)`)'             # code
        r'|(\[(.+?)\]\((.+?)\))'  # link
    )

    last_end = 0
    for match in pattern.finditer(text):
        # Add text before the match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        if match.group(2):  # bold-italic
            run = paragraph.add_run(match.group(2))
            run.font.bold = True
            run.font.italic = True
        elif match.group(4):  # bold
            run = paragraph.add_run(match.group(4))
            run.font.bold = True
        elif match.group(6):  # italic
            run = paragraph.add_run(match.group(6))
            run.font.italic = True
        elif match.group(8):  # code
            run = paragraph.add_run(match.group(8))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif match.group(10):  # link
            link_text = match.group(11)
            run = paragraph.add_run(link_text)
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.font.underline = True

        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def _create_table(doc, rows):
    """Create a table from parsed rows."""
    from docx.shared import Pt, Inches
    from docx.enum.table import WD_ALIGN_VERTICAL

    if not rows:
        return

    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = cell_text.strip()
                # Bold the header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True


def create_from_json(spec, doc):
    """Create document from a JSON specification."""
    from docx.shared import Pt, Inches

    # Set metadata
    if 'metadata' in spec:
        meta = spec['metadata']
        if 'author' in meta:
            doc.core_properties.author = meta['author']
        if 'subject' in meta:
            doc.core_properties.subject = meta['subject']
        if 'title' in meta:
            doc.core_properties.title = meta['title']

    # Add title
    if 'title' in spec:
        doc.add_heading(spec['title'], level=0)

    # Process sections
    for section in spec.get('sections', []):
        # Add heading
        if 'heading' in section:
            level = section.get('level', 1)
            doc.add_heading(section['heading'], level=level)

        # Add content
        if 'content' in section:
            for para_text in section['content'].split('\n'):
                if para_text.strip():
                    doc.add_paragraph(para_text)

        # Add table
        if 'table' in section:
            table_spec = section['table']
            headers = table_spec.get('headers', [])
            rows = table_spec.get('rows', [])

            all_rows = [headers] + rows if headers else rows
            if all_rows:
                _create_table(doc, all_rows)

        # Add list
        if 'list' in section:
            list_type = section.get('list_type', 'bullet')
            style = 'List Bullet' if list_type == 'bullet' else 'List Number'
            for item in section['list']:
                doc.add_paragraph(item, style=style)


def add_header_footer(doc, header_text=None, footer_text=None):
    """Add header and/or footer to the document."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for section in doc.sections:
        if header_text:
            header = section.header
            header.is_linked_to_previous = False
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.text = header_text
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.italic = True

        if footer_text:
            footer = section.footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

            if '{page}' in footer_text:
                # Add page number field
                parts = footer_text.split('{page}')
                p.clear()

                if parts[0]:
                    run = p.add_run(parts[0])
                    run.font.size = Pt(9)

                # Add page number field code
                from lxml import etree
                W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                run_elem = etree.SubElement(p._element, f'{{{W_NS}}}r')
                rpr = etree.SubElement(run_elem, f'{{{W_NS}}}rPr')
                sz = etree.SubElement(rpr, f'{{{W_NS}}}sz')
                sz.set(f'{{{W_NS}}}val', '18')  # 9pt
                fld_char_begin = etree.SubElement(run_elem, f'{{{W_NS}}}fldChar')
                fld_char_begin.set(f'{{{W_NS}}}fldCharType', 'begin')

                run_elem2 = etree.SubElement(p._element, f'{{{W_NS}}}r')
                instr = etree.SubElement(run_elem2, f'{{{W_NS}}}instrText')
                instr.text = ' PAGE '
                instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

                run_elem3 = etree.SubElement(p._element, f'{{{W_NS}}}r')
                fld_char_end = etree.SubElement(run_elem3, f'{{{W_NS}}}fldChar')
                fld_char_end.set(f'{{{W_NS}}}fldCharType', 'end')

                if len(parts) > 1 and parts[1]:
                    run = p.add_run(parts[1])
                    run.font.size = Pt(9)
            else:
                p.text = footer_text
                for run in p.runs:
                    run.font.size = Pt(9)

            p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    parser = argparse.ArgumentParser(description="Create a new Word document")
    parser.add_argument("-i", "--input", default=None,
                        help="Input file (Markdown, text, or JSON)")
    parser.add_argument("--text", default=None,
                        help="Direct text content (alternative to -i)")
    parser.add_argument("-o", "--output", required=True,
                        help="Output .docx file path")
    parser.add_argument("--template", default=None,
                        help="Template .docx for styling")
    parser.add_argument("--header", default=None,
                        help="Header text for all pages")
    parser.add_argument("--footer", default=None,
                        help="Footer text (use {page} for page number)")
    parser.add_argument("--title", default=None,
                        help="Document title (for text input)")

    args = parser.parse_args()

    if not args.input and not args.text:
        print("Error: Provide -i (input file) or --text (direct text)", file=sys.stderr)
        sys.exit(1)

    if not ensure_dependencies():
        sys.exit(1)

    from docx import Document

    # Create document (from template or blank)
    if args.template:
        template_path = os.path.expanduser(args.template)
        if not os.path.exists(template_path):
            print(f"Error: Template not found: {template_path}", file=sys.stderr)
            sys.exit(1)
        doc = Document(template_path)
    else:
        doc = Document()

    # Read input content
    if args.input:
        input_path = os.path.expanduser(args.input)
        input_path = os.path.abspath(input_path)

        if not os.path.exists(input_path):
            print(f"Error: Input file not found: {input_path}", file=sys.stderr)
            sys.exit(1)

        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()

        input_type = detect_input_type(input_path)
    else:
        content = args.text
        input_type = 'text'

    # Create content based on input type
    if input_type == 'json':
        try:
            spec = json.loads(content)
        except json.JSONDecodeError as e:
            print(f"Error: Invalid JSON: {e}", file=sys.stderr)
            sys.exit(1)
        create_from_json(spec, doc)
    elif input_type == 'markdown':
        create_from_markdown(content, doc)
    else:
        create_from_text(content, doc, title=args.title)

    # Add header/footer
    if args.header or args.footer:
        add_header_footer(doc, header_text=args.header, footer_text=args.footer)

    # Save
    output_path = os.path.expanduser(args.output)
    output_path = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)

    file_size = os.path.getsize(output_path)
    print(f"Document created: {output_path} ({file_size:,} bytes)")


if __name__ == "__main__":
    main()
